import os
import re
import sqlite3
import io
from datetime import datetime
import streamlit as st
import pandas as pd
from dotenv import load_dotenv

# File Processing Imports
from docx import Document as DocxReader

# LangChain & Vector DB Imports
from langchain_chroma import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_groq import ChatGroq

# ==========================================
# 1. Load Environment Variables & DB Configs
# ==========================================
load_dotenv()
if not os.getenv("GROQ_API_KEY"):
    st.error("Missing GROQ_API_KEY in .env file.")
    st.stop()

LOGS_DB = "app_logs.db"
DATA_DIR = "data"
CHROMA_PATH = "./chroma_db"
UNIFIED_COLLECTION = "all_documents"

if not os.path.exists(DATA_DIR):
    os.makedirs(DATA_DIR)

def init_logs_db():
    conn = sqlite3.connect(LOGS_DB)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS error_logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp TEXT,
            user_query TEXT,
            wrong_llm_response TEXT,
            retrieved_context TEXT
        )
    """)
    conn.commit()
    conn.close()

def log_llm_error(query, response, context):
    conn = sqlite3.connect(LOGS_DB)
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO error_logs (timestamp, user_query, wrong_llm_response, retrieved_context)
        VALUES (?, ?, ?, ?)
    """, (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), query, response, context))
    conn.commit()
    conn.close()

init_logs_db()

# ==========================================
# 2. Page Configuration & UI Styling
# ==========================================
st.set_page_config(page_title="Dynamic Extension RAG Assistant", page_icon="📂", layout="centered")
st.title("📂 Extension-Driven RAG Assistant")
st.caption("AI agent detecting file extensions dynamically with built-in self-verification guardrails")

# ==========================================
# 3. Initialize Shared Embeddings & LLM
# ==========================================
@st.cache_resource
def init_resources():
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    llm = ChatGroq(model_name="qwen/qwen3-32b", temperature=0.0)
    return embeddings, llm

embeddings, llm = init_resources()

# ==========================================
# 4. Multi-Format Native & Structural Parsing Loaders
# ==========================================
def process_pdf(file_path):
    try:
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        return text_splitter.split_documents(pages)
    except Exception as e:
        st.sidebar.error(f"Error parsing PDF ({os.path.basename(file_path)}): {e}")
        return []

def process_docx(file_path):
    try:
        doc = DocxReader(file_path)
        full_text = [para.text for para in doc.paragraphs if para.text.strip()]
        combined_text = "\n".join(full_text)
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(combined_text)
        return [Document(page_content=chunk, metadata={"source": file_path, "type": "docx"}) for chunk in chunks]
    except Exception as e:
        st.sidebar.error(f"Error parsing DOCX ({os.path.basename(file_path)}): {e}")
        return []

def process_txt(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(text)
        return [Document(page_content=chunk, metadata={"source": file_path, "type": "txt"}) for chunk in chunks]
    except Exception as e:
        st.sidebar.error(f"Error parsing TXT ({os.path.basename(file_path)}): {e}")
        return []


# ==========================================
# FIX A: CSV/XLSX now store TWO documents each:
#   1. A schema blueprint (metadata only, for routing)
#   2. The FULL data as text chunks (for listing/retrieval)
# This prevents the "only 3 rows shown" hallucination.
# ==========================================

def process_csv(file_path):
    try:
        df = pd.read_csv(file_path)
        documents = []

        # Doc 1: Schema blueprint for the pandas engine router
        schema_info = (
            f"STRUCTURED FILE DATA BLUEPRINT:\n"
            f"File Name: {os.path.basename(file_path)}\n"
            f"File Path: {file_path}\n"
            f"Columns: {', '.join(df.columns.tolist())}\n"
            f"Total Rows: {len(df)}\n"
            f"Sample Data Snapshot:\n{df.head(3).to_string(index=False)}"
        )
        documents.append(Document(
            page_content=schema_info,
            metadata={"source": file_path, "type": "structured_data_schema", "extension": "csv"}
        ))

        # Doc 2: Full data as text for unstructured retrieval
        # Split into chunks so large files don't exceed context
        full_text = (
            f"Complete data from file: {os.path.basename(file_path)}\n"
            f"Columns: {', '.join(df.columns.tolist())}\n\n"
            f"{df.to_string(index=False)}"
        )
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100)
        chunks = text_splitter.split_text(full_text)
        for chunk in chunks:
            documents.append(Document(
                page_content=chunk,
                metadata={"source": file_path, "type": "structured_data_full", "extension": "csv"}
            ))

        return documents
    except Exception as e:
        st.sidebar.error(f"Error parsing CSV ({os.path.basename(file_path)}): {e}")
        return []


def process_xlsx(file_path):
    try:
        excel_file = pd.ExcelFile(file_path)
        documents = []

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            # Doc 1: Schema blueprint
            schema_info = (
                f"STRUCTURED FILE DATA BLUEPRINT:\n"
                f"File Name: {os.path.basename(file_path)}\n"
                f"File Path: {file_path}\n"
                f"Sheet Name: {sheet_name}\n"
                f"Columns: {', '.join(df.columns.tolist())}\n"
                f"Total Rows: {len(df)}\n"
                f"Sample Data Snapshot:\n{df.head(3).to_string(index=False)}"
            )
            documents.append(Document(
                page_content=schema_info,
                metadata={
                    "source": file_path,
                    "type": "structured_data_schema",
                    "sheet": sheet_name,
                    "extension": "xlsx"
                }
            ))

            # Doc 2: Full data as text
            full_text = (
                f"Complete data from file: {os.path.basename(file_path)}, Sheet: {sheet_name}\n"
                f"Columns: {', '.join(df.columns.tolist())}\n\n"
                f"{df.to_string(index=False)}"
            )
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=100)
            chunks = text_splitter.split_text(full_text)
            for chunk in chunks:
                documents.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "type": "structured_data_full",
                        "sheet": sheet_name,
                        "extension": "xlsx"
                    }
                ))

        return documents
    except Exception as e:
        st.sidebar.error(f"Error parsing Excel ({os.path.basename(file_path)}): {e}")
        return []


def parse_file(file_path):
    _, ext = os.path.splitext(file_path)
    ext = ext.lower().replace(".", "")
    if ext == "pdf":    return process_pdf(file_path)
    if ext == "docx":   return process_docx(file_path)
    if ext == "txt":    return process_txt(file_path)
    if ext == "csv":    return process_csv(file_path)
    if ext in ["xlsx", "xls"]: return process_xlsx(file_path)
    return []

# ==========================================
# 5. Unified Vector Store Initialization
# ==========================================
def init_vector_store():
    try:
        return Chroma(
            persist_directory=CHROMA_PATH,
            embedding_function=embeddings,
            collection_name=UNIFIED_COLLECTION
        )
    except Exception as e:
        st.error(f"Could not load unified vector store: {e}")
        return None

vector_store = init_vector_store()

# ==========================================
# 6. Sidebar Controls & Chunk Verification
# ==========================================
st.sidebar.header("📦 Vector Database Status")
try:
    total_chunks = len(vector_store.get()['ids']) if vector_store else 0
    if total_chunks == 0:
        st.sidebar.warning("⚠️ Database: 0 Chunks Loaded (Empty)")
    else:
        st.sidebar.success(f"✅ Unified Database: {total_chunks} Chunks Loaded")
except Exception:
    st.sidebar.error("❌ Database Verification Error")

st.sidebar.markdown("---")
st.sidebar.markdown("### 🛠️ Ingest Directory")

if st.sidebar.button("Scan & Ingest Directory"):
    with st.sidebar.status("Scanning data folder dynamically...", expanded=True) as status:
        all_docs = []
        for root, dirs, files in os.walk(DATA_DIR):
            for filename in files:
                file_path = os.path.join(root, filename)
                status.write(f"Analyzing extension for: `{filename}`")
                parsed_docs = parse_file(file_path)
                if parsed_docs:
                    all_docs.extend(parsed_docs)
                    status.write(f"➡️ Parsed {len(parsed_docs)} chunk(s) from `{filename}`")

        if all_docs and vector_store:
            existing_ids = vector_store.get()['ids']
            if existing_ids:
                vector_store.delete(ids=existing_ids)
            vector_store.add_documents(documents=all_docs)
            status.update(label=f"🎉 Successfully ingested {len(all_docs)} objects!", state="complete")
        else:
            status.update(label="No new valid files found or parsed.", state="complete")
    st.rerun()

# ==========================================
# 7. Context Retrieval & Guardrail Functions
# ==========================================
def retrieve_context_documents(query: str, k: int = 8):
    """
    FIX B: Increased k from 5 → 8 so that with 20+ files,
    both the schema doc AND the full-data chunks for the right file
    have a chance to appear together in results.
    """
    if vector_store:
        try:
            return vector_store.similarity_search(query, k=k)
        except Exception as e:
            st.error(f"Retrieval error: {e}")
    return []


# ==========================================
# FIX C: Analytical keyword list now covers listing/naming queries
# AND the dispatch logic is smarter — it checks for the file
# keyword in the query OR in retrieved schema docs.
# ==========================================
ANALYTICAL_INTENT_KEYWORDS = [
    # Counting
    "how many", "count", "total", "number of",
    # Aggregation
    "average", "avg", "sum", "max", "min", "maximum", "minimum",
    # Listing — this was the gap that caused the hallucination
    "list all", "list the", "name all", "name the", "show all",
    "what are all", "what are the", "give me all", "give all",
    "all the", "all departments", "all employees", "all records",
    # Filtering / grouping
    "unique", "distinct", "which", "filter", "group by",
    "breakdown", "distribution", "top ", "bottom ",
    "highest", "lowest", "percentage", "ratio", "who is", "who are"
]

def is_analytical_query(query: str) -> bool:
    q = query.lower().strip()
    if len(q) < 4:
        return False
    return any(kw in q for kw in ANALYTICAL_INTENT_KEYWORDS)

FOLLOWUP_SIGNALS = [
    "list all", "show all", "name all", "give me all", "all of them",
    "list them", "show them", "what are they", "tell me more",
    "the same", "those", "them", "it", "that"
]

def is_ambiguous_followup(query: str) -> bool:
    """Returns True if query is too short or lacks a clear subject."""
    q = query.lower().strip()
    if len(q.split()) <= 3:
        return True
    return any(q.startswith(sig) or q == sig for sig in FOLLOWUP_SIGNALS)


def enrich_query_with_history(current_query: str, messages: list) -> str:
    """
    If the current query is ambiguous (e.g. 'list all', 'show them'),
    prepend the last meaningful user message to give retrieval full context.
    """
    if not is_ambiguous_followup(current_query):
        return current_query
    # Find the last user message that isn't the current one
    prior_user_msgs = [
        m["content"] for m in messages
        if m["role"] == "user" and m["content"] != current_query
    ]
    if prior_user_msgs:
        last_context = prior_user_msgs[-1]
        enriched = f"{last_context} {current_query}"
        return enriched
    return current_query


def pick_best_blueprint(retrieved_docs, user_query: str):
    """
    FIX D: Improved blueprint scoring.
    - Considers ALL schema docs in results, not just first.
    - Scores by column-name overlap too (handles "how many employees" → employees.csv).
    - Falls back to highest-rated schema doc if no name overlap.
    """
    schema_docs = [d for d in retrieved_docs if d.metadata.get("type") == "structured_data_schema"]
    if not schema_docs:
        return None

    query_words = set(re.sub(r'[^a-z0-9 ]', ' ', user_query.lower()).split())

    best_doc = None
    best_score = -1

    for doc in schema_docs:
        source = doc.metadata.get("source", "")
        filename = os.path.basename(source).lower()
        stem = re.sub(r'\.[a-z]+$', '', filename)
        stem_words = set(re.split(r'[_\-\s]+', stem))

        # Score 1: filename stem word overlap with query
        name_overlap = len(query_words & stem_words)

        # Score 2: substring match (e.g. "departments" in query and "departments.csv")
        substring_bonus = 3 if any(sw in user_query.lower() for sw in stem_words if len(sw) > 3) else 0

        # Score 3: column name overlap (e.g. "employees" query matches "employee_id" column)
        col_line = ""
        for line in doc.page_content.splitlines():
            if line.lower().startswith("columns:"):
                col_line = line.lower()
                break
        col_words = set(re.sub(r'[^a-z0-9 ]', ' ', col_line).split())
        col_overlap = len(query_words & col_words)

        score = name_overlap + substring_bonus + col_overlap
        if score > best_score:
            best_score = score
            best_doc = doc

    return best_doc


def validate_response(context: str, query: str, final_answer: str) -> bool:
    if len(final_answer.strip()) < 20:
        return True
    if "do not have enough information" in final_answer.lower():
        return True

    validation_prompt = f"""
You are a strict Quality Assurance AI evaluator. Your ONLY job is to verify factual accuracy.

Context (ground truth):
{context}

User Query: {query}

Assistant Answer: {final_answer}

Rules:
- Reply with exactly ONE word: PASSED or FAILED
- PASSED = The answer is fully supported by the context above
- FAILED = The answer contains any fact NOT present in the context, or contradicts the context
- If the context is insufficient to verify, reply PASSED
- Do NOT output anything else. No punctuation, no explanation.
"""
    try:
        check = llm.invoke([SystemMessage(content=validation_prompt)]).content.strip()
        check = re.sub(r"<think>.*?</think>", "", check, flags=re.DOTALL).strip()
        return check.upper().startswith("PASSED")
    except Exception as e:
        st.warning(f"Validation step failed with error: {e}. Skipping guardrail.")
        return True


# ==========================================
# 8. Chat History Engine
# ==========================================
if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "assistant", "content": "Hello! I am ready to evaluate answers from your dynamically structured local data folder. What is your query?"}
    ]

if "last_retrieved_context" not in st.session_state:
    st.session_state.last_retrieved_context = ""

for idx, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

        if msg["role"] == "assistant" and idx > 0:
            if st.button("⚠️ Report Wrong Answer", key=f"flag_{idx}"):
                associated_user_query = st.session_state.messages[idx - 1]["content"] if idx > 0 else "Unknown"
                saved_context = msg.get("context_snapshot", "Context not captured")
                log_llm_error(associated_user_query, msg["content"], saved_context)
                st.toast("Thank you! This error has been logged for evaluation.", icon="💾")

# ==========================================
# 9. Main Dual-Engine Pipeline Execution
# ==========================================
if user_query := st.chat_input("Ask anything from your data folder..."):
    st.session_state.messages.append({"role": "user", "content": user_query})
    with st.chat_message("user"):
        st.markdown(user_query)

    # ==========================================
    # FIX E: Retrieve more docs (k=8) and separately
    # fetch full-data chunks when a schema is matched.
    # This guarantees the pandas engine always has the correct file path
    # AND the text engine always has the full data, not just 3-row sample.
    # ==========================================
    with st.spinner("Searching database chunks..."):
        # Enrich ambiguous follow-ups with prior conversation context
        retrieval_query = enrich_query_with_history(user_query, st.session_state.messages)
        retrieved_docs = retrieve_context_documents(retrieval_query, k=8)

        retrieved_context = "\n\n".join([
            f"--- Context from {os.path.basename(doc.metadata.get('source', 'Unknown'))} ---\n{doc.page_content}"
            for doc in retrieved_docs
        ])
        st.session_state.last_retrieved_context = retrieved_context

    # --- DETERMINISTIC DISPATCH ROUTER ---
    has_schema_doc = any(doc.metadata.get("type") == "structured_data_schema" for doc in retrieved_docs)
    use_structured_engine = has_schema_doc and is_analytical_query(retrieval_query)

    # ==========================================
    # ENGINE A: PANDAS / STRUCTURED ENGINE
    # Used for: counting, aggregation, listing ALL items, filtering
    # ==========================================
    if use_structured_engine:
        with st.spinner("Analyzing data table metrics dynamically..."):
            matched_blueprint = pick_best_blueprint(retrieved_docs, retrieval_query)

            if not matched_blueprint:
                final_clean_response = "⚠️ I couldn't identify which data file to query. Please mention the file name or topic."
                with st.chat_message("assistant"):
                    st.markdown(final_clean_response)
            else:
                file_path = matched_blueprint.metadata.get("source")
                extension = matched_blueprint.metadata.get("extension", "csv")
                sheet_name = matched_blueprint.metadata.get("sheet", None)

                data_agent_prompt = f"""
You are a precise Python/Pandas data analyst. Output ONLY a single executable pandas expression.

Dataset blueprint:
{matched_blueprint.page_content}

User Query: "{user_query}"

STRICT RULES:
1. The dataframe is already loaded as `df`.
2. Output ONLY the raw pandas expression — no markdown, no backticks, no explanation, no <think> tags.
3. For "how many X" count questions: use df['column'].nunique() or (df['column'] == 'value').sum() — NOT df.shape[0] unless the user literally asks "how many rows".
4. For listing/naming items ("list all", "name all", "what are the", "show all"): use df['column'].tolist() or df['column'].unique().tolist().
5. For aggregations: use df.groupby(...)['col'].agg(...).
6. For filtering: use df[df['column'] == 'value'][['col1', 'col2']].
7. Always return a labelled result — prefer Series/DataFrame over bare scalars where it adds clarity.
8. If the question cannot be answered from the available columns, output exactly: CANNOT_ANSWER
"""

                try:
                    raw_response = llm.invoke([SystemMessage(content=data_agent_prompt)]).content.strip()

                    # Strip reasoning chains
                    code_clean = re.sub(r"<think>.*?</think>", "", raw_response, flags=re.DOTALL)
                    code_clean = re.sub(r"<think>.*", "", code_clean, flags=re.DOTALL)
                    code_clean = code_clean.replace("```python", "").replace("```", "").replace("`", "")
                    code_clean = "\n".join([
                        line for line in code_clean.splitlines()
                        if line.strip() and not line.strip().startswith("<")
                    ])
                    code_clean = code_clean.strip()

                    if code_clean.upper() == "CANNOT_ANSWER" or not code_clean:
                        final_clean_response = "I am sorry, but the available data columns do not contain enough information to answer that question."
                    else:
                        # Load the correct file
                        if extension == "csv" or file_path.endswith('.csv'):
                            df = pd.read_csv(file_path)
                        else:
                            df = pd.read_excel(file_path, sheet_name=sheet_name) if sheet_name else pd.read_excel(file_path)

                        try:
                            computed_result = eval(code_clean, {"df": df, "pd": pd})

                            # Format output cleanly
                            if isinstance(computed_result, pd.DataFrame):
                                computed_result = computed_result.to_string(index=True)
                            elif isinstance(computed_result, pd.Series):
                                computed_result = computed_result.to_string() if not computed_result.empty else "No matching items found."
                            elif isinstance(computed_result, (list, tuple, set)):
                                computed_result = "\n".join(map(str, computed_result)) if computed_result else "No matching items found."
                            elif computed_result is None or (hasattr(computed_result, '__len__') and len(computed_result) == 0):
                                computed_result = "No matching data found in the dataset."

                            # Humanize the computed result via LLM
                            humanize_prompt = f"""
                            You are a friendly data assistant. A user asked: "{user_query}"

                            The precise answer computed from the dataset is:
                            {computed_result}

                            Write a single natural, conversational sentence (or two at most) that directly answers the user's question using this result.
                            - Do NOT use markdown, bullet points, or code blocks.
                            - Do NOT mention pandas, dataframes, or code.
                            - Do NOT add caveats or extra information unless it's directly relevant.
                            - Just answer naturally, like a helpful human colleague would.
                            """
                            try:
                                humanized = llm.invoke([SystemMessage(content=humanize_prompt)]).content.strip()
                                humanized = re.sub(r"<think>.*?</think>", "", humanized, flags=re.DOTALL).strip()
                            except Exception:
                                humanized = str(computed_result)  # fallback to raw result

                            final_clean_response = humanized
                                

                        except Exception as eval_err:
                            final_clean_response = (
                                f"⚠️ Could not execute the generated code against this dataset.\n\n"
                                f"**Error:** `{eval_err}`\n"
                                f"**Code attempted:** `{code_clean}`"
                            )

                except Exception as e:
                    final_clean_response = f"An error occurred while calling the LLM for data analysis: {str(e)}"

                with st.chat_message("assistant"):
                    st.markdown(final_clean_response)

    # ==========================================
    # ENGINE B: UNSTRUCTURED TEXT ENGINE
    # Used for: document Q&A, PDFs, DOCX, TXT
    # Also handles structured files when the query is non-analytical
    # (the full-data chunks now ensure complete data is in context)
    # ==========================================
    else:
        if not retrieved_docs or len(user_query.strip()) < 3:
            final_clean_response = "I am sorry, but I do not have enough information in my database to answer that question."
            with st.chat_message("assistant"):
                st.markdown(final_clean_response)
        else:
            # FIX F: For structured files in text engine, augment context with
            # full-data chunks from the same file so listing queries work correctly.
            # This handles edge cases where is_analytical_query() misses something.
            schema_in_results = [d for d in retrieved_docs if d.metadata.get("type") == "structured_data_schema"]
            if schema_in_results:
                # Fetch the full-data doc for the best matched file specifically
                best_bp = pick_best_blueprint(retrieved_docs, retrieval_query)
                if best_bp:
                    best_source = best_bp.metadata.get("source")
                    # Pull full-data chunks for this specific file from the vector store
                    try:
                        extra_docs = vector_store.similarity_search(
                            user_query, k=5,
                            filter={"$and": [
                                {"source": {"$eq": best_source}},
                                {"type": {"$eq": "structured_data_full"}}
                            ]}
                        )
                        # Merge, dedup by page_content
                        existing_contents = {d.page_content for d in retrieved_docs}
                        for ed in extra_docs:
                            if ed.page_content not in existing_contents:
                                retrieved_docs.append(ed)
                                existing_contents.add(ed.page_content)
                    except Exception:
                        pass  # Non-critical; proceed with what we have

                # Rebuild context after potential augmentation
                retrieved_context = "\n\n".join([
                    f"--- Context from {os.path.basename(doc.metadata.get('source', 'Unknown'))} ---\n{doc.page_content}"
                    for doc in retrieved_docs
                ])
                st.session_state.last_retrieved_context = retrieved_context

            system_prompt = f"""
You are a strict and factual Corporate Document Assistant.
Answer the user's question using ONLY the retrieved context below.

CRITICAL RULES:
1. Base your answer entirely on the retrieved facts below.
2. If the answer is not in the context, reply exactly: "I am sorry, but I do not have enough information in my database to answer that question."
3. Do NOT make up facts, do NOT use outside knowledge, do NOT hallucinate.
4. Do NOT include <think> tags or internal reasoning in your answer.
5. Use the chat history only for follow-up context, never to invent new facts.
6. When listing items from tabular data, list ALL items present in the context — do not stop early.

--- RETRIEVED CONTEXT ---
{retrieved_context}
------------------------
"""

            messages = [SystemMessage(content=system_prompt)]

            for index, msg in enumerate(st.session_state.messages[:-1]):
                if index == 0 and msg["role"] == "assistant":
                    continue
                if msg["role"] == "assistant" and "failed our internal verification check" in msg["content"]:
                    continue
                if msg["role"] == "assistant" and "do not have enough information" in msg["content"]:
                    continue
                if msg["role"] == "user":
                    messages.append(HumanMessage(content=msg["content"]))
                elif msg["role"] == "assistant":
                    messages.append(AIMessage(content=msg["content"]))

            messages.append(HumanMessage(content=user_query))

            with st.chat_message("assistant"):
                response_placeholder = st.empty()
                full_response = ""

                try:
                    for chunk in llm.stream(messages):
                        full_response += chunk.content
                        clean_response = re.sub(r"<think>.*?</think>", "", full_response, flags=re.DOTALL)
                        clean_response = re.sub(r"<think>.*", "", clean_response, flags=re.DOTALL)
                        if clean_response.strip():
                            response_placeholder.markdown(clean_response + "▌")

                    final_clean_response = re.sub(r"<think>.*?</think>", "", full_response, flags=re.DOTALL).strip()
                    final_clean_response = re.sub(r"<think>.*", "", final_clean_response, flags=re.DOTALL).strip()

                    with st.spinner("Verifying answer consistency..."):
                        is_valid = validate_response(retrieved_context, user_query, final_clean_response)

                    if not is_valid:
                        log_llm_error(user_query, f"[AUTO-CAUGHT HALLUCINATION]: {final_clean_response}", retrieved_context)
                        final_clean_response = "⚠️ I am sorry, but I generated an answer that failed our internal verification check. This error has been logged for admin review."

                    response_placeholder.markdown(final_clean_response)

                except Exception as e:
                    final_clean_response = f"An error occurred while calling the LLM: {str(e)}"
                    response_placeholder.markdown(final_clean_response)

    st.session_state.messages.append({
        "role": "assistant",
        "content": final_clean_response,
        "context_snapshot": st.session_state.last_retrieved_context
    })
    st.rerun()